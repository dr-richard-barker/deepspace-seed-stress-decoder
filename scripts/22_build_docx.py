#!/usr/bin/env python
"""Build a submission-ready .docx of the npj manuscript with F1-F6 embedded (python-docx)."""
import os, re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT=r"C:\Users\drric\Downloads\nmf_seed_decoder"
MD=os.path.join(ROOT,"report","manuscript_npj_microgravity.md")
NPJ=os.path.join(ROOT,"report","figures_npj")
OUT=os.path.join(ROOT,"report","manuscript_npj_microgravity.docx")

doc=Document()
s=doc.sections[0]
s.page_width=Inches(8.5); s.page_height=Inches(11)
for m in ("top_margin","bottom_margin","left_margin","right_margin"): setattr(s,m,Inches(1))
normal=doc.styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11)

TOKEN=re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`|\[.+?\]\(.+?\))")
def add_runs(p,text):
    text=text.replace("\\*","*").replace("\\","")
    for tok in TOKEN.split(text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"): r=p.add_run(tok[2:-2]); r.bold=True
        elif tok.startswith("*") and tok.endswith("*"): r=p.add_run(tok[1:-1]); r.italic=True
        elif tok.startswith("`") and tok.endswith("`"):
            r=p.add_run(tok[1:-1]); r.font.name="Consolas"; r.font.size=Pt(9.5)
        elif tok.startswith("[") and "](" in tok:
            r=p.add_run(re.match(r"\[(.+?)\]",tok).group(1))
        else: p.add_run(tok)

lines=open(MD,encoding="utf-8").read().split("\n")
for ln in lines:
    t=ln.rstrip()
    if not t.strip(): continue
    if t.startswith("# "):       doc.add_heading(t[2:].strip(),0)
    elif t.startswith("## "):    doc.add_heading(t[3:].strip(),1)
    elif t.startswith("### "):   doc.add_heading(t[4:].strip(),2)
    elif t.strip()=="---":       continue
    elif re.match(r"^\d+\.\s",t):
        p=doc.add_paragraph(style="List Number"); add_runs(p,re.sub(r"^\d+\.\s","",t))
    elif t.startswith("> "):
        p=doc.add_paragraph(); add_runs(p,t[2:]); p.runs and setattr(p.runs[0],"italic",True)
    elif t.startswith("- "):
        p=doc.add_paragraph(style="List Bullet"); add_runs(p,t[2:])
    else:
        p=doc.add_paragraph(); add_runs(p,t)

# Figures
doc.add_page_break(); doc.add_heading("Figures",1)
plates=[("F1_concept_workflow.png","Figure 1. Two-tool framework: signature → DSRS / GSAD → DeepSpace atlas."),
        ("F2_stress_library.png","Figure 2. DSRS stress reference library — 22 contrasts (5 families) → seed programs."),
        ("F3_seed_reference_embryo_lineage.png","Figure 3. Seed reference validation — developing-embryo state → germinating cell type recovers textbook lineages."),
        ("F4_gsad_susceptibility.png","Figure 4. GSAD — stressor → germinating-seed cell-type susceptibility."),
        ("F5_deepspace_atlas.png","Figure 5. DeepSpace seed-susceptibility atlas — cell-type × stressor family + convergence (radicle-tip hotspot)."),
        ("F6_convergence_model.png","Figure 6. Radicle growth-point: deep-space multi-stressor convergence model + falsification tests.")]
for fn,cap in plates:
    p=os.path.join(NPJ,fn)
    if os.path.exists(p):
        doc.add_picture(p,width=Inches(6.3))
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(); r=c.add_run(cap); r.italic=True; r.font.size=Pt(9)

doc.save(OUT)

# patch python-docx w:zoom (missing required w:percent) so the file validates
import zipfile, shutil
with zipfile.ZipFile(OUT,"r") as z:
    names=z.namelist(); data={n:z.read(n) for n in names}
st=data["word/settings.xml"].decode("utf-8")
st=re.sub(r"<w:zoom\s*/>", '<w:zoom w:percent="100"/>', st)
st=re.sub(r'(<w:zoom\b(?![^>]*w:percent)[^>]*)/>', r'\1 w:percent="100"/>', st)
data["word/settings.xml"]=st.encode("utf-8")
tmp=OUT+".tmp"
with zipfile.ZipFile(tmp,"w",zipfile.ZIP_DEFLATED) as z:
    for n in names: z.writestr(n,data[n])
shutil.move(tmp,OUT)
print("DOCX:",OUT,"| size KB",round(os.path.getsize(OUT)/1024),"| zoom patched")
